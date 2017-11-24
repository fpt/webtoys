#!/usr/bin/env python
# coding:utf-8


from argparse import ArgumentParser, FileType
import docx
from docx.shared import Inches
from itertools import chain

from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.document import CT_Body
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


class DocReader(object):
    """Factory"""

    def __init__(self):
        pass

    def process(self, fp):

        # extract w:p and w:tbl somehow
        def iter_block_items(doc):
            assert isinstance(doc, Document)

            d = {}
            for p in doc.paragraphs:
                d[p._element] = p
            for t in doc.tables:
                d[t._element] = t

            es = doc._element.xpath('//w:tbl[not(ancestor::w:tbl)] | //w:p[not(ancestor::w:tbl)]')
            for idx, e in enumerate(es):
                if e in d.keys():
                    o = d[e]
                    if isinstance(o, Paragraph):
                        yield do_paragraph(o)
                    elif isinstance(o, Table):
                        for t in do_table(o):
                            yield t
                else:
                    pass
                    # print(e)

        def do_paragraph(par):
            return par.text.strip()

        def do_table(tbl):
            for row in tbl.rows:
                r = []
                for c in row.cells:
                    r.extend(list(do_cell(c)))
                yield '\t'.join(r)

        def do_cell(c):
            yield c.text.strip()
            for t in c.tables:
                return do_table(t)

        doc = docx.Document(fp)
        # https://github.com/python-openxml/python-docx/issues/40#issuecomment-42096998
        return iter_block_items(doc)


def main():
    parser = ArgumentParser()
    parser.add_argument('infile', nargs='?', type=FileType('rb'))
    args = parser.parse_args()

    dr = DocReader()
    text = '\n'.join(list(dr.process(args.infile)))
    print(text)

if __name__ == '__main__':
    main()
