#!/usr/bin/env python
# coding:utf-8


from argparse import ArgumentParser, FileType
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image, ImageDraw

class BaseWriter(object):
    """Factory"""

    def __init__(self):
        pass

    def add_heading(self, str, level=None):
        raise "notimpl"

    def add_paragraph(self, str):
        raise "notimpl"

    def add_image(self, img):
        pass

    def add_page_break(self):
        pass

    def close(self):
        raise "notimpl"


class TxtWriter(BaseWriter):
    """Factory"""

    def __init__(self, fp):
        self.fp = fp

    def add_heading(self, str, level=None):
        self.fp.write(str + "\n")

    def add_paragraph(self, str):
        self.fp.write(str + "\n")

    def close(self):
        pass


class DocWriter(BaseWriter):
    """Factory"""

    def __init__(self, fp):
        self.fp = fp
        self.document = Document()

    def add_heading(self, str, level=None):
        self.document.add_heading(str, level=level)

    def add_paragraph(self, str):
        self.document.add_paragraph(str)

    def add_image(self, img):
        img_bytes = BytesIO()
        img.save(img_bytes, format="png")
        img_bytes.seek(0)
        self.document.add_picture(img_bytes)

    def add_page_break(self):
        self.document.add_page_break()

    def close(self):
        self.document.save(self.fp)


def main():
    parser = ArgumentParser()
    args = parser.parse_args()
    dw = DocWriter()
    dw.save()

if __name__ == '__main__':
    main()
